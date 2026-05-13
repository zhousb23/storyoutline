# -*- coding: utf-8 -*-
"""结果管理模块 — 项目状态读写、JSON 存储、Word/TXT 导出。"""

import json
import os
from dataclasses import dataclass, field

from src.chapter_splitter import Chapter


# ============================================================
# 项目状态管理
# ============================================================

@dataclass
class ProjectState:
    """项目元数据和进度。"""

    name: str = ""
    source_file: str = ""  # 原始文件路径
    chapters: list[dict] = field(default_factory=list)  # [{index, title, content, status}]
    step1_results: list[dict] = field(default_factory=list)
    step2_results: list[dict] = field(default_factory=list)
    step3_result: dict | None = None
    character_analysis: str = ""  # 人物分析结果
    writing_style: str = ""  # 写作风格分析结果
    writing_advice: str = ""  # 写作建议 + 段落摘抄
    current_stage: str = "idle"  # idle / step1 / step2 / step3 / done
    analyzed_count: int = 0  # 已完成第一步分析的章节数
    section_size: int = 50  # 每阶段章节数


class ResultManager:
    """管理项目数据的读写和导出。"""

    def __init__(self, project_dir: str):
        self.project_dir = project_dir
        os.makedirs(project_dir, exist_ok=True)

    # ---- 状态读写 ----

    def save_state(self, state: ProjectState) -> None:
        """保存项目状态到 JSON 文件。"""
        data = {
            "name": state.name,
            "source_file": state.source_file,
            "chapters": state.chapters,
            "current_stage": state.current_stage,
            "analyzed_count": state.analyzed_count,
            "section_size": state.section_size,
        }
        self._write_json("project.json", data)

    def load_state(self) -> ProjectState:
        """读取项目状态。"""
        data = self._read_json("project.json")
        if not data:
            return ProjectState()
        return ProjectState(
            name=data.get("name", ""),
            source_file=data.get("source_file", ""),
            chapters=data.get("chapters", []),
            current_stage=data.get("current_stage", "idle"),
            analyzed_count=data.get("analyzed_count", 0),
            section_size=data.get("section_size", 50),
        )

    # ---- 分析结果读写 ----

    def save_step1_results(self, results: list[dict]) -> None:
        """保存第一步（章节大纲）结果。"""
        self._write_json("step1_results.json", results)

    def load_step1_results(self) -> list[dict]:
        """读取第一步结果。"""
        return self._read_json("step1_results.json") or []

    def save_step2_results(self, results: list[dict]) -> None:
        """保存第二步（阶段总结）结果。"""
        self._write_json("step2_results.json", results)

    def load_step2_results(self) -> list[dict]:
        """读取第二步结果。"""
        return self._read_json("step2_results.json") or []

    def save_step3_result(self, result: dict) -> None:
        """保存第三步（全书故事线）结果。"""
        self._write_json("step3_result.json", result)

    def load_step3_result(self) -> dict | None:
        """读取第三步结果。"""
        return self._read_json("step3_result.json")

    # ---- 导出 ----

    def export_to_docx(self, filepath: str, project_name: str = "故事大纲分析") -> None:
        """导出为格式化的 Word 文档。"""
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 设置默认字体
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Microsoft YaHei"
        font.size = Pt(11)

        # -- 封面标题 --
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(project_name)
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x3B, 0x7D, 0xD8)
        doc.add_paragraph()

        # -- 第一章节大纲 --
        doc.add_heading("一、章节大纲", level=1)
        self._export_step1_to_docx(doc)

        # -- 第二阶段性总结 --
        self._export_step2_to_docx(doc)

        # -- 第三全书故事线 --
        self._export_step3_to_docx(doc)

        # -- 人物分析 --
        self._export_character_to_docx(doc)

        # -- 风格分析 --
        self._export_style_to_docx(doc)

        # -- 写作建议 --
        self._export_writing_advice_to_docx(doc)

        doc.save(filepath)

    def export_to_txt(self, filepath: str, project_name: str = "故事大纲分析") -> None:
        """导出为纯文本文件。"""
        lines = [
            f"{'=' * 50}",
            f"  {project_name}",
            f"{'=' * 50}",
            "",
        ]
        self._export_sections_to_text(lines)

        with open(filepath, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))

    # ---- 内部方法 ----

    def _filepath(self, filename: str) -> str:
        return os.path.join(self.project_dir, filename)

    def _write_json(self, filename: str, data) -> None:
        with open(self._filepath(filename), "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

    def _read_json(self, filename: str):
        filepath = self._filepath(filename)
        if not os.path.isfile(filepath):
            return None
        with open(filepath, "r", encoding="utf-8") as f:
            return json.load(f)

    def _export_step1_to_docx(self, doc) -> None:
        """将第一步结果写入 Word 文档。"""
        results = self.load_step1_results()
        if not results:
            doc.add_paragraph("（暂无数据）")
            return

        for r in results:
            chapter_title = f"第{r['chapter_index']}章 {r.get('chapter_title', '')}"
            doc.add_heading(chapter_title, level=2)

            fields = [
                ("激励事件", r.get("inciting_incident", "")),
                ("进展纠葛", r.get("progressive_complications", "")),
                ("危机", r.get("crisis", "")),
                ("高潮", r.get("climax", "")),
                ("结局", r.get("resolution", "")),
            ]
            for label, content in fields:
                if content:
                    p = doc.add_paragraph()
                    run_label = p.add_run(f"**{label}**：")
                    run_label.font.bold = True
                    p.add_run(content)

    def _export_step2_to_docx(self, doc) -> None:
        """将第二步结果写入 Word 文档。"""
        results = self.load_step2_results()
        if not results:
            return

        doc.add_heading("二、阶段性总结", level=1)
        for r in results:
            section_title = f"阶段{r['section_index']}（第{r['chapter_range'][0]}-{r['chapter_range'][1]}章）"
            doc.add_heading(section_title, level=2)
            doc.add_paragraph(r.get("summary", ""))

    def _export_step3_to_docx(self, doc) -> None:
        """将第三步结果写入 Word 文档。"""
        result = self.load_step3_result()
        if not result:
            return

        doc.add_heading("三、全书故事线", level=1)
        doc.add_paragraph(result.get("overall_storyline", ""))

    def _export_character_to_docx(self, doc) -> None:
        """将人物分析写入 Word 文档。"""
        data = self._read_json("character_analysis.json")
        if not data:
            return
        doc.add_heading("四、人物小传", level=1)
        doc.add_paragraph(data.get("result", ""))

    def _export_style_to_docx(self, doc) -> None:
        """将风格分析写入 Word 文档。"""
        data = self._read_json("style_analysis.json")
        if not data:
            return
        doc.add_heading("五、语言风格分析", level=1)
        doc.add_paragraph(data.get("result", ""))

    def _export_writing_advice_to_docx(self, doc) -> None:
        """将写作建议写入 Word 文档。"""
        data = self._read_json("writing_advice.json")
        if not data:
            return
        doc.add_heading("六、写作建议与优秀段落摘抄", level=1)
        doc.add_paragraph(data.get("result", ""))

    def _export_sections_to_text(self, lines: list) -> None:
        """将各步结果追加到文本行列表。"""
        # 第一步
        results1 = self.load_step1_results()
        if results1:
            lines.append("一、章节大纲")
            lines.append("-" * 40)
            for r in results1:
                lines.append(f"\n第{r['chapter_index']}章 {r.get('chapter_title', '')}")
                lines.append(f"  激励事件：{r.get('inciting_incident', '')}")
                lines.append(f"  进展纠葛：{r.get('progressive_complications', '')}")
                lines.append(f"  危机：{r.get('crisis', '')}")
                lines.append(f"  高潮：{r.get('climax', '')}")
                lines.append(f"  结局：{r.get('resolution', '')}")
            lines.append("")

        # 第二步
        results2 = self.load_step2_results()
        if results2:
            lines.append("二、阶段性总结")
            lines.append("-" * 40)
            for r in results2:
                lines.append(f"\n阶段{r['section_index']}（第{r['chapter_range'][0]}-{r['chapter_range'][1]}章）")
                lines.append(r.get("summary", ""))
            lines.append("")

        # 第三步
        result3 = self.load_step3_result()
        if result3:
            lines.append("三、全书故事线")
            lines.append("-" * 40)
            lines.append(result3.get("overall_storyline", ""))

        # 人物分析
        char_data = self._read_json("character_analysis.json")
        if char_data:
            lines.append("\n四、人物小传")
            lines.append("-" * 40)
            lines.append(char_data.get("result", ""))
            lines.append("")

        # 风格分析
        style_data = self._read_json("style_analysis.json")
        if style_data:
            lines.append("五、语言风格分析")
            lines.append("-" * 40)
            lines.append(style_data.get("result", ""))
            lines.append("")

        # 写作建议
        advice_data = self._read_json("writing_advice.json")
        if advice_data:
            lines.append("六、写作建议与优秀段落摘抄")
            lines.append("-" * 40)
            lines.append(advice_data.get("result", ""))
